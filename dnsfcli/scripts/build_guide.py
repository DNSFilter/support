"""DNSFilter brand-locked guide builder.

Renders the same DNSFilter-branded guide content into four formats:
  - PDF                (final delivery, brand-perfect)
  - Markdown (.md)     (GitHub-flavored, README-style sharing)
  - DOCX (.docx)       (editable handoff, comments, reviewers)
  - Confluence (.xhtml)(storage format, push to Confluence via REST API
                        or paste into a "view source" editor)

Run in place from this assets directory, or copy alongside its `fonts/`
and `logos/` siblings. The script self-resolves those folders via __file__.

Usage:
    1. Edit the CONTENT block below: cover fields, sections, callouts, etc.
    2. (Optional) Switch COVER from "minimal" to "feature" for the
       blue header band cover.
    3. Set OUT_BASENAME and OUT_DIR for your output files.
    4. (Optional) Restrict FORMATS to a subset.
    5. Run: python build_guide.py

Brand rules baked in (do not override):
    palette = white #ffffff, Cyber Black #000000, DNSFilter Blue #3427fd,
              Threat Magenta #f306ae (warnings only), Secure Sky Blue
              #00c8fd (success only), light gray #f5f5f5, mid gray #666666
    typography = Montserrat (display, headers, labels) + Inter (body)
    logo top-left on cover. Identity never on Threat Magenta.

Voice and structure are flexible. Palette and typography are not.

Requires: pip install weasyprint beautifulsoup4 python-docx cairosvg
"""

import io
import re
from html import escape as html_escape
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from cairosvg import svg2png
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from weasyprint import CSS, HTML

ASSETS = Path(__file__).resolve().parent
FONTS = ASSETS / "fonts"
LOGOS = ASSETS / "logos"

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------

OUT_DIR = Path(".")
OUT_BASENAME = "dnsfcli-reference"
FORMATS = ["pdf", "md", "docx", "confluence"]  # build all by default

# Cover styles:
#   "minimal" - white background, logo top-left, large title (default for
#               SOPs, internal docs, technical walkthroughs).
#   "feature" - DNSFilter Blue header band with white logo, white below
#               with large title (default for lead magnets and external
#               marketing guides).
COVER = "feature"

# ---------------------------------------------------------------------------
# CONTENT
# Edit the strings below. The HTML body is the canonical source for all
# four formats. Keep the tags and class names; replace the text.
# ---------------------------------------------------------------------------

COVER_FIELDS = {
    "doc_type": "Reference · 2026",
    "eyebrow": "Command-line interface for the DNSFilter API",
    "title": "dnsfcli<br>Reference Guide",
    "subtitle": (
        "Complete command reference, tested examples, CSV bulk-import "
        "workflows, and output flag usage for every endpoint in the "
        "DNSFilter API."
    ),
    "byline": [
        ("Tool version", "0.1.0"),
        ("Updated", "2026-07-13"),
        ("Audience", "Developers & Administrators"),
    ],
}

BODY = """
<section class="content">

<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Overview</h2>

<p><strong>dnsfcli</strong> is a Python command-line tool that wraps the entire DNSFilter REST API. Every endpoint, every write field, and every output mode is available directly from the terminal — no browser required.</p>

<div class="callout">
  <div class="callout-title">Command structure</div>
  <div class="callout-body">All commands follow a single three-part pattern: <code>dnsfcli.py [endpoint] [function] [--param value …]</code>. The endpoint is the resource group (e.g. <code>networks</code>, <code>policies</code>), the function is the operation (e.g. <code>list</code>, <code>create</code>, <code>update</code>), and flags supply the values.</div>
</div>

<p>The tool routes to 242 operations across 36 resource groups, all defined in a single endpoint registry. Mistyped endpoint or function names exit non-zero with a &ldquo;did you mean&rdquo; suggestion.</p>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Installation &amp; Requirements</h2>

<div class="step">
  <div class="step-num">1</div>
  <div class="step-content">
    <div class="step-title">Install from the repository</div>
    <div class="step-body">pip installs straight from GitHub — re-run the same command to upgrade. To work from source, the tool lives in the <code>dnsfcli/</code> directory of the public <code>DNSFilter/support</code> repository and uses a standard <code>src</code>-layout Python package.</div>
  </div>
</div>

<div class="step">
  <div class="step-num">2</div>
  <div class="step-content">
    <div class="step-title">Install dependencies</div>
    <div class="step-body">Python 3.11 or later is required. Install with pip from the project root.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / INSTALL</div>
  <div class="prompt-body">pip install "git+https://github.com/DNSFilter/support.git#subdirectory=dnsfcli"
# or, from a local clone:
git clone https://github.com/DNSFilter/support
cd support/dnsfcli
pip install -e .
# installs: typer, click, httpx, keyring, rich</div>
</div>

<div class="step">
  <div class="step-num">3</div>
  <div class="step-content">
    <div class="step-title">Run directly or as a module</div>
    <div class="step-body">Both invocation styles are equivalent. The entry-point script <code>dnsfcli.py</code> in the project root is the recommended way to run it during development.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / INVOCATION OPTIONS</div>
  <div class="prompt-body">python dnsfcli.py [endpoint] [function] [flags]
python -m dnsfcli  [endpoint] [function] [flags]
dnsfcli            [endpoint] [function] [flags]   # after pip install</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Authentication</h2>

<p>The tool stores credentials in the OS keychain (macOS Keychain, Windows Credential Manager, Linux Secret Service). Credentials are read automatically on every command — no manual token passing required once set up.</p>

<div class="step">
  <div class="step-num">1</div>
  <div class="step-content">
    <div class="step-title">Store your API token</div>
    <div class="step-body">Run <code>auth setup</code> and enter your DNSFilter JWT or API key when prompted. Optionally set a default org ID so it pre-fills templates and org-scoped calls.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / AUTH SETUP</div>
  <div class="prompt-body">python dnsfcli.py auth setup
# DNSFilter API key: ••••••••••••••

python dnsfcli.py auth setup --org-id 802315</div>
</div>

<div class="step">
  <div class="step-num">2</div>
  <div class="step-content">
    <div class="step-title">Verify the credentials work</div>
    <div class="step-body"><code>auth verify</code> makes a live call to <code>/v1/organizations</code> and confirms the token is valid before you run anything else.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / AUTH VERIFY</div>
  <div class="prompt-body">python dnsfcli.py auth verify
✓ Credentials are valid.

python dnsfcli.py auth show
  api_key    eyJhbG...k3Qw
  org_id     802315
  base_url   https://api.dnsfilter.com</div>
</div>

<div class="callout">
  <div class="callout-title">Environment variable override</div>
  <div class="callout-body">Set <code>DNSF_API_KEY</code> to override the keychain token for a single session. Pass <code>--api-key &lt;TOKEN&gt;</code> on any command to override for a single call.</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Global Flags</h2>

<p>These flags work on every command and can be placed anywhere in the argument list — before the endpoint, between endpoint and function, or after all other flags.</p>

<table style="width:100%; border-collapse:collapse; font-size:0.78em; margin:0.8em 0;">
<tbody>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Output &amp; format</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--raw, -r</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Print the raw JSON response instead of the formatted table.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--json / --jsonl</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">JSON (or one object per line) on stdout. Automatic when output is piped.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--output FMT</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Unified format switch: table, json, jsonl, raw, csv, or none.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--columns a,b,c</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Limit table/CSV output to the named columns.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--preset NAME</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Apply a named column preset from config (column_presets.NAME).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--format TMPL</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Render each row through a template — both "{name} ({id})" and "{{.name}}" styles work.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--format-preset N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Apply a named format template from config (format_presets.NAME).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--bundle NAME</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Apply a named command bundle from config (columns + filter + sort + format).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--pick FIELD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Print a single field, one value per line — ideal for piping.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--jq PATH</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Extract a dotted path (data.0.name) from the response before output.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--to-csv FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Write the response to FILE as CSV (use - for stdout).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--to-json FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Write the response to FILE as JSON.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--to-markdown FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Write the response to FILE as a Markdown table (use - for stdout).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--tee FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Save a copy of everything printed to FILE.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--quiet, -q</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Suppress status chatter; result data still prints.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--no-color</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Disable ANSI colour output (NO_COLOR env also honoured).</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Filtering, sorting &amp; aggregation</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--filter EXPR</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Keep matching rows: field=value, field!=value, field~substr, field&gt;N, &gt;=, &lt;, &lt;=.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--filter-mode or</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">OR logic across multiple --filter flags (default: AND).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--grep REGEX</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Keep rows where any field matches the regex.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--unique FIELD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Drop rows with a duplicate value in FIELD.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--sort FIELD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Sort by field; prefix with - for descending (--sort -created_at).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--limit N / --last N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Return at most N results / the last N results.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--count</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Print the result count only.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--count-by FIELD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Frequency table of FIELD values.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--group-by FIELD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Group results by FIELD value.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--sum / --avg / --min / --max F</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Aggregate a numeric field and print one value.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--select a,b / --exclude a,b</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Keep only (or remove) the named fields from every item.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--not-null F / --is-null F</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Keep rows where FIELD is (or is not) null.</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Pagination</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--all</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Fetch every page of paginated results.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--page N / --page-size N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Request a specific page / page size.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--max-pages N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Cap the number of pages fetched by --all.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--paginate-until E</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Stop --all pagination once any item matches filter expression E.</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Batch &amp; CSV input</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--from-csv FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Read input rows from FILE — one API call per row (use - for stdin).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--template</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Print a blank CSV import template for this command and exit. No auth needed.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--skip-rows N / --max-rows N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Resume an interrupted batch / cap rows processed.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--batch-report FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Write a JSON run summary (per-row outcomes, counts) to FILE.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--on-error MODE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">stop or continue when a batch row fails.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--errors-to-csv FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Write failed rows to FILE for later retry (--retry-errors-csv).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--confirm-each</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Prompt before each batch row.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--validate-only</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Validate the CSV and exit without calling the API.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--dry-run / --plan</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Show what would be sent without making any API calls.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--yes, -y</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Skip confirmation prompts.</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Transformation</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--add-field K=V</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Inject a static field into every result item.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--transform F=EXPR</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Compute a new field from a restricted expression (ratio=blocked/total).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--map FIELD=OP</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Transform a field value: upper, lower, strip, title, truncate:N.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--join EP:LK=RK</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Client-side join: fetch endpoint EP, match local key LK to remote key RK.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--rename A=B</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Rename field A to B in the output.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--flatten</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Flatten nested objects into dotted keys.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--strip-nulls</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Remove null-valued keys from every item.</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Multi-organization</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--each-org</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Run the command once per organization on the account.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--org-csv FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Load the --each-org organization list from a CSV file.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--org-filter REGEX</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Restrict --each-org to organizations whose name matches.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--max-orgs N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Cap the number of organizations processed.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--parallel-orgs</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Process organizations concurrently (--org-concurrency N).</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Watch, monitor &amp; CI</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--watch N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Re-run the command every N seconds (Ctrl-C to stop).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--watch-until EXPR</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Stop watching once any result matches the filter expression.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--watch-diff</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Show only what changed between watch ticks.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--alert EXPR</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Ring the terminal bell and print a banner when a result matches.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--fail-on-empty</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Exit 1 when the result list is empty (messages go to stderr).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--fail-on-pattern E</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Exit 1 when any result row matches filter expression E.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--exec CMD</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Run a shell command per result row with {field} substitution (values are shell-quoted).</td></tr>
<tr><td colspan="2" style="padding:8px 9px 4px; font-weight:bold; font-size:0.95em; color:#3427fd; border-bottom:2px solid #3427fd;">Connection &amp; authentication</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--api-key TOKEN</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Override the keychain token for this call only (prefer DNSF_API_KEY).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--org-id ID</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Override the stored organization ID for this call only.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--profile NAME</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Use a named credential profile.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--timeout N / --connect-timeout N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Read / connect timeouts in seconds.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--rate N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Client-side request-per-second cap.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--retry N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Retry attempts for failed batch rows.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--header K=V</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Add a custom request header (scrubbed from history logs).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--proxy URL</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Route requests through a proxy (scrubbed from history logs).</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--env-file FILE</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Load KEY=VALUE pairs into the environment before running.</td></tr>
<tr><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5; font-family:monospace; white-space:nowrap;">--cache-ttl N</td><td style="padding:4px 9px; border-bottom:1px solid #e5e5e5;">Serve identical GETs from a local cache for N seconds.</td></tr>
</tbody>
</table>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Discovery</h2>

<p>Two built-in commands let you explore the full endpoint catalogue without leaving the terminal.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / LIST ALL ENDPOINTS</div>
  <div class="prompt-body">python dnsfcli.py endpoints
# Endpoint                  Functions
# ─────────────────────────────────────────────────────────
# agent-local-users         bulk-delete, delete, list, show, update …
# api-keys                  create, delete, list, revoke, show
# networks                  bulk-create, counts, create, delete, list …
# policies                  add-blacklist-domain, create, delete, list …
# traffic-reports           qps, query-logs, top-domains, total-requests …
# … (36 groups total)</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / FUNCTIONS FOR ONE ENDPOINT</div>
  <div class="prompt-body">python dnsfcli.py endpoints policies
# add-allowed-application   add-blacklist-category
# add-blacklist-domain      add-whitelist-domain
# application               bulk-add-allowlist
# create                    delete
# list                      list-all
# permissive-mode           remove-blacklist-domain
# set-permissive-mode       show
# update</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Core Resource Operations</h2>

<p>Every major resource follows the same five-function pattern. The examples below use <strong>policies</strong> but the same functions apply to networks, organizations, block-pages, ip-addresses, mac-addresses, and scheduled-policies.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / LIST</div>
  <div class="prompt-body">python dnsfcli.py policies list

# ┏━━━━━━━━━┳━━━━━━━━━━┳━━━━━━━━━━━━━━━━━━━━━━┳━━━━━━━━━━━━━━━━━━┓
# ┃ id      ┃ type     ┃ attributes           ┃ relationships    ┃
# ┡━━━━━━━━━╇━━━━━━━━━━╇━━━━━━━━━━━━━━━━━━━━━━╇━━━━━━━━━━━━━━━━━━┩
# │ 331207 │ policies │ name=big bada        │ organization=…   │
# │         │          │ blocklist,           │                  │
# │         │          │ organization_id=…    │                  │
# │ 285109  │ policies │ name=Block Adult     │ organization=…   │
# │         │          │ Content, …           │                  │
# └─────────┴──────────┴──────────────────────┴──────────────────┘

python dnsfcli.py policies list --raw | jq '.data[].attributes.name'</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / SHOW A SINGLE RESOURCE</div>
  <div class="prompt-body">python dnsfcli.py policies show --id 285109

# ╭──────────────────────── policies show ─────────────────────────╮
# │  id                              285109                        │
# │  attributes.name                 Block Adult Content           │
# │  attributes.organization_id      802315                        │
# │  attributes.blacklist_categories 56, 69, 2, 38                │
# │  attributes.google_safesearch    no                            │
# │  attributes.youtube_restricted   yes                           │
# │  attributes.allow_unknown_domains yes                          │
# ╰─────────────────────────────────────────────────────────────────╯</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / CREATE</div>
  <div class="prompt-body">python dnsfcli.py policies create \
  --name "Guest WiFi" \
  --organization_id 802315 \
  --allow_unknown_domains true \
  --youtube_restricted true \
  --youtube_restricted_level strict \
  --google_safesearch true \
  --bing_safe_search true</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / UPDATE</div>
  <div class="prompt-body">python dnsfcli.py policies update --id 285109 \
  --name "Block Adult Content v2" \
  --interstitial true</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / DELETE</div>
  <div class="prompt-body">python dnsfcli.py policies delete --id 285109</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Networks</h2>

<div class="callout">
  <div class="callout-title">Policy assignment uses an array</div>
  <div class="callout-body">Networks use <code>policy_ids</code> (a JSON array), not a single <code>policy_id</code>. Pass it as a JSON array string: <code>--policy_ids '["285109","331207"]'</code>.</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / NETWORKS — KEY OPERATIONS</div>
  <div class="prompt-body"># List with org filter
python dnsfcli.py networks list --organization_id 802315

# Create with policy assignment
python dnsfcli.py networks create \
  --name "Branch Office" \
  --organization_id 802315 \
  --policy_ids '["285109"]' \
  --physical_address "123 Main St, Denver CO"

# Network counts
python dnsfcli.py networks counts

# Geographic data
python dnsfcli.py networks geo

# LAN IP management
python dnsfcli.py networks lan-ips --id 736401
python dnsfcli.py networks lan-ip-update --id 736401 --lan_ip_id 42 \
  --name "Reception Desk"

# Subnets (use 'from' and 'to', not cidr)
python dnsfcli.py networks subnets-create --id 736401 \
  --name "Sales Floor" --from "10.0.1.0" --to "10.0.1.255" \
  --policy_id 285109

# Secret key rotation
python dnsfcli.py networks secret-key-create --id 736401</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Policy Domain &amp; Category Management</h2>

<p>Individual domains and categories are added or removed from a policy with targeted action functions. Application filtering uses the application <em>name</em> string, not an ID.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / POLICY DOMAIN &amp; CATEGORY ACTIONS</div>
  <div class="prompt-body"># Block a domain
python dnsfcli.py policies add-blacklist-domain \
  --id 285109 --domain "malware.example.com" \
  --note "Flagged by threat intel"

# Allow a domain
python dnsfcli.py policies add-whitelist-domain \
  --id 285109 --domain "internal.corp.com"

# Block a category (Adult Content = 2)
python dnsfcli.py policies add-blacklist-category \
  --id 285109 --category_id 2

# Block an application by name
python dnsfcli.py policies add-blocked-application \
  --id 285109 --name "TikTok"

# Remove a domain from blocklist
python dnsfcli.py policies remove-blacklist-domain \
  --id 285109 --domain "malware.example.com"

# Bulk add domains to multiple policies at once
python dnsfcli.py policies bulk-add-blocklist \
  --policy_ids '["285109","331207"]' \
  --domains '["evil.com","malware.net"]'</div>
</div>

<div class="callout warn">
  <div class="callout-title">Application names are case-sensitive</div>
  <div class="callout-body">The <code>add-blocked-application</code> and <code>add-allowed-application</code> functions take the application <code>name</code> field exactly as it appears in <code>python dnsfcli.py applications list</code>. Mismatched casing returns a 404.</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Organizations &amp; Users</h2>

<div class="prompt">
  <div class="prompt-label">SHELL / ORGANIZATIONS</div>
  <div class="prompt-body"># List all organizations
python dnsfcli.py organizations list

# Show one organization
python dnsfcli.py organizations show --id 802315

# Create an organization
python dnsfcli.py organizations create \
  --name "New Client Corp" \
  --billing_contact_email "billing@newclient.com" \
  --sku "professional"

# Organization settings
python dnsfcli.py organizations settings

# Promote to MSP
python dnsfcli.py organizations promote-to-msp</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / ORGANIZATION USERS</div>
  <div class="prompt-body"># List users in an org
python dnsfcli.py organizations users-list --organization_id 802315

# Add a user
python dnsfcli.py organizations users-create \
  --organization_id 802315 \
  --email "newuser@company.com" \
  --first_name "Jane" --last_name "Smith" \
  --role "administrator"

# Update a user's role
python dnsfcli.py organizations users-update \
  --organization_id 802315 --id 42618 \
  --role "read_only"

# Remove a user
python dnsfcli.py organizations users-delete \
  --organization_id 802315 --id 42618</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>IP &amp; MAC Addresses</h2>

<div class="callout">
  <div class="callout-title">Field name correction</div>
  <div class="callout-body">The API field is <code>address</code>, not <code>ip_address</code> or <code>mac_address</code>. Both IP and MAC address endpoints use the same <code>address</code> parameter name.</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / IP &amp; MAC ADDRESSES</div>
  <div class="prompt-body"># Add a static IP to a network
python dnsfcli.py ip-addresses create \
  --address "203.0.113.5" \
  --organization_id 802315 \
  --network_id 736401

# Add a MAC address
python dnsfcli.py mac-addresses create \
  --organization_id 802315 \
  --address "AA:BB:CC:DD:EE:FF" \
  --filter_value "Reception Printer" \
  --policy_id 285109

# Show my current IP
python dnsfcli.py ip-addresses myip</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Roaming Agents (User Agents)</h2>

<div class="prompt">
  <div class="prompt-label">SHELL / USER AGENTS</div>
  <div class="prompt-body"># List all agents with filters
python dnsfcli.py user-agents list \
  --organization_id 802315 \
  --network_id 736401

# Agent counts
python dnsfcli.py user-agents counts

# Update an agent (note: display name field is 'friendly_name')
python dnsfcli.py user-agents update \
  --id "9b2f6c04-5a1e-4d37-8c2a-f0e1d2c3b4a5" \
  --friendly_name "Finance-Laptop" \
  --policy_id 285109 \
  --tags '["managed","finance"]'

# Bulk update multiple agents at once
python dnsfcli.py user-agent-bulk-updates create \
  --ids '["9b2f6c04-5a1e-4d37-8c2a-f0e1d2c3b4a5"]' \
  --policy_id 285109

# CSV export of all agents
python dnsfcli.py user-agents csv > agents.csv

# Uninstall PIN
python dnsfcli.py user-agents uninstall-pin</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Traffic Reports</h2>

<p>All 51 traffic report endpoints are GET-only and require <code>--start_date</code> and <code>--end_date</code> in <code>YYYY-MM-DD</code> format. Results can be scoped to an org or network with optional filters.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / TRAFFIC REPORTS</div>
  <div class="prompt-body"># Total requests in January
python dnsfcli.py traffic-reports total-requests \
  --start_date 2025-01-01 --end_date 2025-01-31

# Top blocked domains last 30 days
python dnsfcli.py traffic-reports top-domains \
  --start_date 2025-01-01 --end_date 2025-01-31 \
  --organization_id 802315

# Threat summary
python dnsfcli.py traffic-reports total-threats \
  --start_date 2025-01-01 --end_date 2025-01-31

# Query logs (export raw DNS events)
python dnsfcli.py traffic-reports query-logs \
  --start_date 2025-01-01 --end_date 2025-01-31 \
  --csv query-logs-jan.csv

# QPS for active agents
python dnsfcli.py traffic-reports qps-active-agents \
  --start_date 2025-01-01 --end_date 2025-01-31</div>
</div>

<div class="callout warn">
  <div class="callout-title">Some report endpoints enforce a maximum time window</div>
  <div class="callout-body"><code>total-client-stats</code> and a few QPS endpoints accept only short time ranges (roughly 20 minutes). Use them for real-time monitoring, not historical analysis.</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Domain Lookups &amp; Classification</h2>

<div class="prompt">
  <div class="prompt-label">SHELL / DOMAIN LOOKUPS</div>
  <div class="prompt-body"># Bulk classify multiple domains
python dnsfcli.py domains bulk-lookup \
  --domains "google.com,facebook.com,malware.example"

# Look up a domain for a specific user session
python dnsfcli.py domains user-lookup --domain "google.com"

# Suggest a domain for threat review
python dnsfcli.py domains suggest-threat \
  --domain "suspicious.example.com" \
  --reason "Flagged in phishing email campaign"</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>API Key Management</h2>

<div class="prompt">
  <div class="prompt-label">SHELL / API KEYS</div>
  <div class="prompt-body"># List all API keys
python dnsfcli.py api-keys list

# Create a new key (expiry must be within 1 year)
python dnsfcli.py api-keys create \
  --name "CI Pipeline Key" \
  --expiry "2027-05-31"

# Revoke a key immediately
python dnsfcli.py api-keys revoke --id 99

# Delete a key permanently
python dnsfcli.py api-keys delete --id 99</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Scheduled Reports &amp; Policies</h2>

<div class="prompt">
  <div class="prompt-label">SHELL / SCHEDULED REPORTS</div>
  <div class="prompt-body"># Create a weekly threat report
python dnsfcli.py scheduled-reports create \
  --organization_id 802315 \
  --frequency "weekly" \
  --day_of_week "1" \
  --include_threat_summary true \
  --include_content_category_summary true \
  --send_to_dashboard_users true

# Preview a report immediately
python dnsfcli.py scheduled-reports preview-create \
  --organization_id 802315 \
  --include_threat_summary true</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / SCHEDULED POLICIES</div>
  <div class="prompt-body"># Create a time-based policy schedule
python dnsfcli.py scheduled-policies create \
  --name "School Hours" \
  --organization_id 802315 \
  --policy_ids '["285109"]' \
  --timezone "America/Denver"</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>v2 Endpoints</h2>

<p>The v2 API surface covers Cyber Sight, CSV exports, UI settings, and VPN dictionary data. Commands follow the same structure using the <code>v2-</code> endpoint prefix.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / V2 ENDPOINTS</div>
  <div class="prompt-body"># UI settings for current user
python dnsfcli.py v2-current-user ui-settings
python dnsfcli.py v2-current-user ui-settings-update \
  --theme_mode "dark"

# Cyber Sight activity types reference
python dnsfcli.py v2-dictionary cyber-sight-activity-types

# VPN settings state types
python dnsfcli.py v2-dictionary vpn-settings-state-types

# Cyber Sight CSV export
python dnsfcli.py v2-cyber-sight csv-export \
  --organization_ids '["802315"]' \
  --threats_only true \
  --start_at "2025-01-01T00:00:00Z" \
  --end_at "2025-01-31T23:59:59Z"

# Agent local user counts (v2)
python dnsfcli.py v2-agent-local-users counts

# Networks CSV export (v2)
python dnsfcli.py v2-networks csv-export \
  --organization_ids '["802315"]'</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Output Modes</h2>

<p>By default every command renders a human-readable rich table. Three flags change the output format.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / OUTPUT — TABLE (DEFAULT)</div>
  <div class="prompt-body">python dnsfcli.py policies show --id 285109
# ╭──────────────────── policies show ─────────────────────╮
# │  attributes.name                Block Adult Content     │
# │  attributes.organization_id     802315                  │
# │  attributes.youtube_restricted  yes                     │
# │  attributes.google_safesearch   no                      │
# ╰─────────────────────────────────────────────────────────╯</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / OUTPUT — RAW JSON (--raw)</div>
  <div class="prompt-body">python dnsfcli.py policies show --id 285109 --raw
# {
#   "id": 285109,
#   "type": "policies",
#   "attributes": {
#     "name": "Block Adult Content",
#     "organization_id": 802315,
#     ...
#   }
# }

# Pipe directly into jq
python dnsfcli.py policies list --raw | jq '.[].attributes.name'</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / OUTPUT — CSV FILE (--csv)</div>
  <div class="prompt-body"># Save list results to a CSV file
python dnsfcli.py policies list --csv policies.csv
# ✓ Wrote 8 rows to policies.csv

# Traffic report to CSV for Excel/Sheets
python dnsfcli.py traffic-reports total-requests \
  --start_date 2025-01-01 --end_date 2025-01-31 \
  --csv jan-traffic.csv

# Combine --raw and --csv: raw JSON to screen, CSV saved simultaneously
python dnsfcli.py networks list --raw --csv networks-backup.csv</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Bulk CSV Import (--from-csv)</h2>

<p>Any write operation can accept a CSV file as its input source. The tool validates every row before making a single API call, then processes them sequentially with per-row status output.</p>

<div class="step">
  <div class="step-num">1</div>
  <div class="step-content">
    <div class="step-title">Generate a blank template</div>
    <div class="step-body">The <code>--template</code> flag prints a correctly structured CSV with comment lines documenting required and optional fields. If a default org ID is stored in the keychain, it is pre-filled in the example row.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / GENERATE TEMPLATE</div>
  <div class="prompt-body">python dnsfcli.py policies create --template

# Template : dnsfcli policies create
# Required : name (string), organization_id (integer)
# Optional : allow_unknown_domains (boolean), google_safesearch (boolean) …
name,organization_id,allow_unknown_domains,google_safesearch, …
example text,802315,,,…

# Save directly to file
python dnsfcli.py policies create --template > policies-template.csv</div>
</div>

<div class="step">
  <div class="step-num">2</div>
  <div class="step-content">
    <div class="step-title">Fill in the template</div>
    <div class="step-body">Open the CSV in Excel, Google Sheets, or any editor. Fill in one row per resource. The comment lines (<code>#</code> prefix) can be left in — they are skipped by the importer. Arrays use JSON syntax: <code>["item1","item2"]</code>.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">CSV / FILLED TEMPLATE EXAMPLE</div>
  <div class="prompt-body"># Template : dnsfcli policies create
# Required : name (string), organization_id (integer)
name,organization_id,google_safesearch,youtube_restricted,allow_unknown_domains
Guest WiFi,802315,true,true,true
Employee Standard,802315,true,false,false
Kiosk Restrictive,802315,true,true,false</div>
</div>

<div class="step">
  <div class="step-num">3</div>
  <div class="step-content">
    <div class="step-title">Import with --from-csv</div>
    <div class="step-body">The tool validates all rows first. Any errors are shown with row numbers and the field that failed — and no API calls are made until the entire file is clean.</div>
  </div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / IMPORT</div>
  <div class="prompt-body">python dnsfcli.py policies create --from-csv policies-template.csv

# Processing 3 rows from CSV…
#   Row 1: ✓ (id: 1501234)
#   Row 2: ✓ (id: 1501235)
#   Row 3: ✓ (id: 1501236)
#
# Done: 3/3 rows succeeded</div>
</div>

<div class="callout secure">
  <div class="callout-title">Path params can come from either the CLI or the CSV</div>
  <div class="callout-body">Supply path parameters (like <code>--id</code>) on the command line to apply the same value to every row. Or include an <code>id</code> column in the CSV to target a different resource per row — useful for bulk updates.</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / MIXED CLI + CSV — BULK DOMAIN BLOCK</div>
  <div class="prompt-body"># Block a list of domains against a single policy
# CSV: domain,note
# evil.com,Phishing campaign
# malware.net,C2 infrastructure

python dnsfcli.py policies add-blacklist-domain \
  --id 285109 \
  --from-csv threats.csv

# Processing 2 rows from CSV…
#   Row 1: ✓
#   Row 2: ✓
# Done: 2/2 rows succeeded</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Validation &amp; Error Handling</h2>

<p>Errors are always surfaced cleanly — no Python tracebacks are shown to the user. Three categories of errors may occur.</p>

<div class="prompt">
  <div class="prompt-label">SHELL / ERROR — MISSING PATH PARAMETER</div>
  <div class="prompt-body">python dnsfcli.py policies show
# Error: Required path parameter --id was not provided.
# Path template: /v1/policies/{id}</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / ERROR — CSV STRUCTURAL PROBLEM</div>
  <div class="prompt-body">python dnsfcli.py policies create --from-csv bad.csv
# Error: CSV validation failed for bad.csv:
#   Missing required column(s): name
#     Required columns : name (string), organization_id (integer)
#     Optional columns : allow_unknown_domains (boolean) …
#     Run with --template to generate a blank example CSV
#
#   No API calls were made</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / ERROR — CSV DATA PROBLEM</div>
  <div class="prompt-body">python dnsfcli.py networks create --from-csv networks.csv
# Error: CSV validation failed for networks.csv:
#   Row 3: 'organization_id' -- expected an integer, got 'acme-corp'
#   Row 5: 'name' is required but empty
#
#   2 error(s) found across 2 row(s) -- no API calls were made</div>
</div>

<div class="prompt">
  <div class="prompt-label">SHELL / ERROR — API ERROR</div>
  <div class="prompt-body">python dnsfcli.py policies delete --id 99999
# Error: HTTP 404: Unable to find the object that you requested.
#   error: Unable to find the object that you requested.</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Retry &amp; Backoff Behaviour</h2>

<p>The HTTP client handles transient failures automatically so scripts do not need retry logic of their own.</p>

<table style="width:100%; border-collapse:collapse; font-size:0.85em; margin:1em 0;">
  <thead>
    <tr style="background:#f5f5f5;">
      <th style="padding:6px 10px; text-align:left; border-bottom:2px solid #e5e5e5;">Condition</th>
      <th style="padding:6px 10px; text-align:left; border-bottom:2px solid #e5e5e5;">Behaviour</th>
    </tr>
  </thead>
  <tbody>
    <tr><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">HTTP 429 (rate limited)</td><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">Reads <code>Retry-After</code> header, sleeps that many seconds, retries indefinitely until the server responds with non-429.</td></tr>
    <tr><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">Connection error / timeout</td><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">Exponential backoff from 1 s, capped at 60 s, up to 6 attempts before raising.</td></tr>
    <tr><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">HTTP 5xx server error</td><td style="padding:5px 10px; border-bottom:1px solid #e5e5e5;">Raised immediately as <code>APIError</code> — not retried. Operator-side issues are not transient.</td></tr>
    <tr><td style="padding:5px 10px;">HTTP 4xx (except 429)</td><td style="padding:5px 10px;">Raised immediately with the error message from the response body.</td></tr>
  </tbody>
</table>

<div class="callout">
  <div class="callout-title">Rate limit for the DNSFilter API</div>
  <div class="callout-body">The API enforces 2,000 requests per 300 seconds per organization. Large <code>--from-csv</code> imports against a heavily used account may hit this. The client handles the 429 automatically; no special flag is needed.</div>
</div>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Complete Write Field Reference</h2>

<p>Every write endpoint with its full parameter set. <strong>Bold</strong> fields are required.</p>

<table style="width:100%; border-collapse:collapse; font-size:0.78em; margin:1em 0;">
  <thead>
    <tr style="background:#f5f5f5;">
      <th style="padding:5px 8px; text-align:left; border-bottom:2px solid #e5e5e5;">Endpoint / Function</th>
      <th style="padding:5px 8px; text-align:left; border-bottom:2px solid #e5e5e5;">Required</th>
      <th style="padding:5px 8px; text-align:left; border-bottom:2px solid #e5e5e5;">Optional</th>
    </tr>
  </thead>
  <tbody>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>agent-local-users update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">friendly_name, policy_id, scheduled_policy_id, block_page_id</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>agent-local-users bulk-delete</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">ids (array)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">exclude_ids (array)</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>api-keys create</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">expiry (YYYY-MM-DD, max 1 year)</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>billing update-address</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id (path)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">first_name, last_name, email, company, phone, line1, line2, line3, city, state, state_code, zip, country</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>block-pages create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id, block_org_name, block_email_addr, block_logo_uuid</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>current-user update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">—</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">first_name, last_name, phone</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>ip-addresses create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">address, organization_id, network_id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">dynamic_hostname</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>mac-addresses create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">address, filter_value, policy_id, scheduled_policy_id, block_page_id</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>networks create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name, organization_id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">block_page_id, policy_ids (array), external_id, is_legacy_vpn_active, physical_address, local_domains (array), local_resolvers (array)</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>networks subnets-create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id (network), name, from (start IP), to (end IP)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">policy_id, scheduled_policy_id, block_page_id</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>organizations create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">billing_contact_name, billing_contact_phone, billing_contact_email, address, show_pii_rc_hostnames, unique_id, sku, quantity, gdpr, privacy_mode, enable_cybersight</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>organizations users-create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id (path)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">email, first_name, last_name, phone, role, organization_permission_ids (array), is_include_only_list</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>policies create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name, organization_id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">allow_unknown_domains, google_safesearch, bing_safe_search, duck_duck_go_safe_search, ecosia_safesearch, yandex_safe_search, youtube_restricted, youtube_restricted_level, interstitial, allow_list_only, is_global_policy, policy_ip_id, whitelist_domains (array), blacklist_domains (array), blacklist_categories (array), allow_applications (array), block_applications (array), append_domains, include_relationships</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>policies add/remove domain</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id, domain</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">note, include_relationships</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>policies add/remove category</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id, category_id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">include_relationships</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>policies add/remove application</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id, name (app name string)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">include_relationships</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>policies bulk-add/remove domain lists</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">policy_ids (array), domains (array)</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">—</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>scheduled-policies create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">name</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id, policy_ids (array), timezone</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>scheduled-reports create / update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">—</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">organization_id, frequency, day_of_week, include_threat_summary, include_content_category_summary, content_categories_show_count, send_to_dashboard_users, scheduled_report_recipients (array), selected_sub_orgs (array)</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>user-agent-bulk-updates create</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">—</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">ids (array), exclude_ids (array), network_id, policy_id, scheduled_policy_id, block_page_id, friendly_name, tags (array), release_channels (array)</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>user-agents update</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">id</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">friendly_name, status, network_id, policy_id, scheduled_policy_id, block_page_id, tags (array)</td></tr>
    <tr style="background:#f9f9f9;"><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;"><strong>users change-password</strong></td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">new_password</td><td style="padding:4px 8px; border-bottom:1px solid #e5e5e5;">—</td></tr>
    <tr style="background:#fff;"><td style="padding:4px 8px;"><strong>v2-current-user ui-settings-update</strong></td><td style="padding:4px 8px;">—</td><td style="padding:4px 8px;">disable_license_warnings, user_uuid, theme_mode</td></tr>
  </tbody>
</table>


<!-- ═══════════════════════════════════════════════════════════════════ -->
<h2>Quick Start</h2>
<ol class="checklist">
  <li>Run <code>python dnsfcli.py auth setup</code> and store your API token and org ID.</li>
  <li>Run <code>python dnsfcli.py endpoints</code> to see all 40 resource groups, then <code>python dnsfcli.py [endpoint] --template</code> to get a pre-filled CSV for any write operation.</li>
  <li>Run <code>python dnsfcli.py policies list --csv policies-backup.csv</code> to export a baseline snapshot before making bulk changes.</li>
</ol>

<div class="signoff">
  <div class="signoff-text">From the team at</div>
  <div class="signoff-name">DNSFilter</div>
</div>

</section>
"""

# ---------------------------------------------------------------------------
# BRAND CONSTANTS (do not edit)
# ---------------------------------------------------------------------------

DNSF_BLACK = "#000000"
DNSF_BLUE = "#3427fd"
DNSF_MAGENTA = "#f306ae"
DNSF_SKY = "#00c8fd"
HAIRLINE = "#e5e5e5"
BG_LIGHT = "#f5f5f5"
TEXT_MUTED = "#666666"


def hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def load_logo_svg(name):
    raw = (LOGOS / name).read_text()
    raw = re.sub(r"<\?xml[^?]*\?>\s*", "", raw)
    raw = re.sub(r"<!--.*?-->\s*", "", raw, flags=re.DOTALL)
    return raw


LOGO_DARK_SVG = load_logo_svg("horizontal-dark.svg")
LOGO_WHITE_SVG = load_logo_svg("horizontal-white.svg")


def rasterize_svg(svg_str, width_px=900):
    """Rasterize an inline SVG string to PNG bytes.

    cairosvg requires the SVG to declare a size. If the source uses only
    viewBox (which is fine for browsers and weasyprint), inject explicit
    width/height attributes from the viewBox before handing it off so the
    rasterization for .md/.docx/.confluence outputs does not crash.
    """
    if "<svg" in svg_str and " width=" not in svg_str.split(">", 1)[0]:
        m = re.search(r'viewBox\s*=\s*"([\d.\s\-]+)"', svg_str)
        if m:
            parts = m.group(1).split()
            if len(parts) == 4:
                vb_w, vb_h = parts[2], parts[3]
                svg_str = svg_str.replace(
                    "<svg",
                    f'<svg width="{vb_w}" height="{vb_h}"',
                    1,
                )
    return svg2png(bytestring=svg_str.encode("utf-8"), output_width=width_px)


# ---------------------------------------------------------------------------
# PARSE: HTML body -> structured component list
# Each component is a tuple. Renderers consume this intermediate.
# ---------------------------------------------------------------------------


def _inner_html(node):
    return "".join(str(c) for c in node.children).strip()


def _inner_text(node):
    return node.get_text(" ", strip=True)


def parse_body(body_html):
    soup = BeautifulSoup(body_html, "html.parser")
    out = []
    fig_counter = 0

    root = soup.find("section", class_="content") or soup
    for el in root.find_all(recursive=False):
        if isinstance(el, NavigableString):
            continue
        name = el.name
        classes = el.get("class", []) or []

        if name == "h2":
            out.append(("section", _inner_text(el)))

        elif name == "p":
            out.append(("paragraph", _inner_html(el)))

        elif name == "div" and "callout" in classes:
            variant = "info"
            if "warn" in classes:
                variant = "warn"
            elif "secure" in classes:
                variant = "secure"
            title_node = el.select_one(".callout-title")
            body_node = el.select_one(".callout-body")
            out.append((
                "callout",
                variant,
                _inner_text(title_node) if title_node else "",
                _inner_html(body_node) if body_node else "",
            ))

        elif name == "div" and "step" in classes:
            num_node = el.select_one(".step-num")
            title_node = el.select_one(".step-title")
            body_node = el.select_one(".step-body")
            out.append((
                "step",
                _inner_text(num_node) if num_node else "",
                _inner_text(title_node) if title_node else "",
                _inner_html(body_node) if body_node else "",
            ))

        elif name == "div" and "prompt" in classes:
            label_node = el.select_one(".prompt-label")
            body_node = el.select_one(".prompt-body")
            out.append((
                "prompt",
                _inner_text(label_node) if label_node else "",
                _inner_html(body_node) if body_node else "",
            ))

        elif name == "div" and "tools" in classes:
            tools = []
            for t in el.select(".tool"):
                tools.append((
                    _inner_text(t.select_one(".tool-name")) if t.select_one(".tool-name") else "",
                    _inner_text(t.select_one(".tool-meta")) if t.select_one(".tool-meta") else "",
                    _inner_text(t.select_one(".tool-desc")) if t.select_one(".tool-desc") else "",
                ))
            out.append(("tools", tools))

        elif name == "ol" and "checklist" in classes:
            items = [_inner_html(li) for li in el.find_all("li", recursive=False)]
            out.append(("checklist", items))

        elif name == "div" and "signoff" in classes:
            text_node = el.select_one(".signoff-text")
            name_node = el.select_one(".signoff-name")
            out.append((
                "signoff",
                _inner_text(text_node) if text_node else "",
                _inner_text(name_node) if name_node else "DNSFilter",
            ))

        elif name == "figure" or (name == "div" and "figure" in classes):
            fig_counter += 1
            svg_node = el.find("svg")
            cap_node = el.find("figcaption") or el.select_one(".figcaption")
            svg_str = str(svg_node) if svg_node else ""
            caption = _inner_text(cap_node) if cap_node else ""
            out.append(("figure", fig_counter, svg_str, caption))

    return out


# ---------------------------------------------------------------------------
# Inline HTML conversion helpers (markdown + plain text)
# ---------------------------------------------------------------------------

_TAG_TO_MD = [
    (re.compile(r"<strong>(.*?)</strong>", re.DOTALL | re.IGNORECASE), r"**\1**"),
    (re.compile(r"<b>(.*?)</b>", re.DOTALL | re.IGNORECASE), r"**\1**"),
    (re.compile(r"<em>(.*?)</em>", re.DOTALL | re.IGNORECASE), r"_\1_"),
    (re.compile(r"<i>(.*?)</i>", re.DOTALL | re.IGNORECASE), r"_\1_"),
    (re.compile(r"<code>(.*?)</code>", re.DOTALL | re.IGNORECASE), r"`\1`"),
    (re.compile(r'<a\s+[^>]*href="([^"]+)"[^>]*>(.*?)</a>', re.DOTALL | re.IGNORECASE), r"[\2](\1)"),
    (re.compile(r"<br\s*/?>", re.IGNORECASE), "\n"),
]


def html_to_md_inline(html):
    text = html or ""
    for pat, repl in _TAG_TO_MD:
        text = pat.sub(repl, text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def html_to_plain(html):
    text = re.sub(r"<br\s*/?>", " ", html or "", flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\s+", " ", text).strip()


# ---------------------------------------------------------------------------
# RENDERER: PDF (existing WeasyPrint pipeline, body HTML stays canonical)
# ---------------------------------------------------------------------------


def _cover_html_pdf():
    byline_cols = "\n".join(
        f'<div class="byline-col"><div class="byline-label">{label}</div>'
        f'<div class="byline-value">{value}</div></div>'
        for label, value in COVER_FIELDS["byline"]
    )
    if COVER == "feature":
        return f"""
<section class="cover cover-feature">
  <div class="cover-band">
    <div class="cover-logo">{LOGO_WHITE_SVG}</div>
    <div class="doc-type doc-type-light">{COVER_FIELDS["doc_type"]}</div>
  </div>
  <div class="cover-middle">
    <div class="eyebrow">{COVER_FIELDS["eyebrow"]}</div>
    <h1>{COVER_FIELDS["title"]}</h1>
    <div class="rule"></div>
    <div class="subtitle">{COVER_FIELDS["subtitle"]}</div>
  </div>
  <div class="cover-bottom">{byline_cols}</div>
</section>
"""
    return f"""
<section class="cover cover-minimal">
  <div class="cover-top">
    <div class="cover-logo">{LOGO_DARK_SVG}</div>
    <div class="doc-type">{COVER_FIELDS["doc_type"]}</div>
  </div>
  <div class="cover-middle">
    <div class="eyebrow">{COVER_FIELDS["eyebrow"]}</div>
    <h1>{COVER_FIELDS["title"]}</h1>
    <div class="rule"></div>
    <div class="subtitle">{COVER_FIELDS["subtitle"]}</div>
  </div>
  <div class="cover-bottom">{byline_cols}</div>
</section>
"""


PDF_CSS = f"""
@font-face {{ font-family: 'Montserrat'; src: url('file://{FONTS}/Montserrat-VariableFont_wght.ttf'); font-weight: 100 900; font-style: normal; }}
@font-face {{ font-family: 'Montserrat'; src: url('file://{FONTS}/Montserrat-Italic-VariableFont_wght.ttf'); font-weight: 100 900; font-style: italic; }}
@font-face {{ font-family: 'Inter'; src: url('file://{FONTS}/Inter-VariableFont_opsz,wght.ttf'); font-weight: 100 900; font-style: normal; }}
@font-face {{ font-family: 'Inter'; src: url('file://{FONTS}/Inter-Italic-VariableFont_opsz,wght.ttf'); font-weight: 100 900; font-style: italic; }}

@page {{ size: Letter; margin: 0.4in 0.4in 0.5in 0.4in; background: #ffffff;
  @bottom-center {{ content: counter(page); font-family: 'Inter','Arial',sans-serif; font-size: 9pt; color: {TEXT_MUTED}; }}
}}
@page :first {{ margin: 0; @bottom-center {{ content: ""; }} }}

* {{ box-sizing: border-box; }}
html, body {{ background: #ffffff; }}
body {{ font-family: 'Inter','Arial',sans-serif; color: {DNSF_BLACK}; font-size: 11pt; line-height: 1.6; margin: 0; padding: 0; }}

.cover {{ page: cover; width: 8.5in; height: 11in; position: relative; page-break-after: always; display: flex; flex-direction: column; }}
.cover-minimal {{ background: #ffffff; color: {DNSF_BLACK}; padding: 0.4in; justify-content: space-between; }}
.cover-top {{ display: flex; justify-content: space-between; align-items: center; border-bottom: 0.75pt solid {DNSF_BLACK}; padding-bottom: 14pt; }}
.cover-logo svg {{ height: 28pt; width: auto; }}
.doc-type {{ font-family: 'Inter','Arial',sans-serif; font-style: italic; font-size: 10pt; color: {DNSF_BLACK}; }}
.cover-middle {{ flex-grow: 1; display: flex; flex-direction: column; justify-content: center; max-width: 6.5in; padding: 1in 0; }}
.eyebrow {{ font-family: 'Inter','Arial',sans-serif; font-style: italic; font-size: 13pt; color: {DNSF_BLUE}; margin-bottom: 18pt; }}
.cover h1 {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 64pt; line-height: 1.0; font-weight: 800; margin: 0 0 26pt 0; letter-spacing: -1.5pt; color: {DNSF_BLACK}; }}
.rule {{ width: 0.6in; height: 4pt; background: {DNSF_BLUE}; margin: 0 0 22pt 0; }}
.subtitle {{ font-family: 'Inter','Arial',sans-serif; font-size: 14pt; line-height: 1.5; color: {DNSF_BLACK}; max-width: 5.5in; font-weight: 400; }}
.cover-bottom {{ display: flex; justify-content: flex-start; gap: 0.55in; border-top: 0.75pt solid {DNSF_BLACK}; padding-top: 14pt; }}
.byline-col {{ min-width: 1.6in; }}
.byline-label {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 8pt; letter-spacing: 1pt; text-transform: uppercase; font-weight: 600; color: {TEXT_MUTED}; margin-bottom: 3pt; }}
.byline-value {{ font-family: 'Inter','Arial',sans-serif; font-size: 11pt; color: {DNSF_BLACK}; }}

.cover-feature {{ background: #ffffff; color: {DNSF_BLACK}; padding: 0; }}
.cover-feature .cover-band {{ background: {DNSF_BLUE}; color: #ffffff; padding: 0.5in 0.4in 0.4in 0.4in; display: flex; justify-content: space-between; align-items: center; }}
.cover-feature .cover-logo svg {{ height: 32pt; }}
.cover-feature .doc-type-light {{ color: #ffffff; }}
.cover-feature .cover-middle {{ padding: 0.4in; }}
.cover-feature .cover-bottom {{ padding: 14pt 0.4in 0.4in 0.4in; border-top: 0.75pt solid {DNSF_BLACK}; }}

.content {{ padding-top: 0.1in; }}
h2 {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 20pt; font-weight: 700; color: {DNSF_BLACK}; margin: 30pt 0 12pt 0; padding-bottom: 8pt; border-bottom: 0.75pt solid {DNSF_BLACK}; letter-spacing: -0.3pt; }}
p {{ margin: 0 0 11pt 0; font-family: 'Inter','Arial',sans-serif; color: {DNSF_BLACK}; }}
b, strong {{ color: {DNSF_BLACK}; font-weight: 700; }}
a {{ color: {DNSF_BLUE}; text-decoration: none; border-bottom: 0.5pt solid {DNSF_BLUE}; }}

.callout {{ background: {BG_LIGHT}; border-left: 3pt solid {DNSF_BLUE}; padding: 14pt 18pt; margin: 16pt 0; }}
.callout-title {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 8pt; font-weight: 700; letter-spacing: 1.5pt; text-transform: uppercase; color: {DNSF_BLUE}; margin-bottom: 5pt; }}
.callout-body {{ font-family: 'Inter','Arial',sans-serif; font-size: 11pt; color: {DNSF_BLACK}; line-height: 1.6; }}
.callout.warn {{ border-left-color: {DNSF_MAGENTA}; }}
.callout.warn .callout-title {{ color: {DNSF_MAGENTA}; }}
.callout.secure {{ border-left-color: {DNSF_SKY}; }}
.callout.secure .callout-title {{ color: {DNSF_SKY}; }}

.step {{ display: flex; align-items: flex-start; margin: 0; padding: 12pt 0; border-bottom: 0.5pt solid {HAIRLINE}; }}
.step:last-of-type {{ border-bottom: none; }}
.step-num {{ font-family: 'Montserrat','Arial',sans-serif; width: 28pt; font-weight: 700; font-size: 16pt; color: {DNSF_BLUE}; margin-right: 14pt; flex-shrink: 0; line-height: 1; padding-top: 2pt; }}
.step-content {{ flex: 1; }}
.step-title {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 12pt; font-weight: 600; color: {DNSF_BLACK}; margin-bottom: 3pt; }}
.step-body {{ font-family: 'Inter','Arial',sans-serif; font-size: 11pt; color: {DNSF_BLACK}; line-height: 1.55; }}

.prompt {{ background: {DNSF_BLACK}; color: #ffffff; padding: 16pt 20pt; margin: 16pt 0; border-left: 3pt solid {DNSF_BLUE}; }}
.prompt-label {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 8pt; font-weight: 700; letter-spacing: 2pt; color: {DNSF_BLUE}; margin-bottom: 8pt; text-transform: uppercase; }}
.prompt-body {{ font-family: 'Inter','Arial',sans-serif; font-size: 11pt; line-height: 1.6; color: #ffffff; font-style: italic; }}

.tools {{ margin: 12pt 0; }}
.tool {{ border-top: 0.5pt solid {HAIRLINE}; padding: 12pt 0; }}
.tool:last-child {{ border-bottom: 0.5pt solid {HAIRLINE}; }}
.tool-name {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 13pt; font-weight: 600; color: {DNSF_BLACK}; }}
.tool-meta {{ font-family: 'Inter','Arial',sans-serif; font-size: 10pt; color: {DNSF_BLUE}; font-style: italic; margin: 2pt 0 4pt 0; }}
.tool-desc {{ font-family: 'Inter','Arial',sans-serif; font-size: 11pt; color: {DNSF_BLACK}; }}

.checklist {{ padding-left: 0; list-style: none; counter-reset: item; margin: 12pt 0; }}
.checklist li {{ counter-increment: item; padding: 10pt 0 10pt 36pt; position: relative; font-family: 'Inter','Arial',sans-serif; font-size: 11pt; border-bottom: 0.5pt solid {HAIRLINE}; color: {DNSF_BLACK}; }}
.checklist li:first-child {{ border-top: 0.5pt solid {HAIRLINE}; }}
.checklist li::before {{ content: counter(item, decimal-leading-zero); position: absolute; left: 0; top: 10pt; font-family: 'Montserrat','Arial',sans-serif; color: {DNSF_BLUE}; font-weight: 700; font-size: 11pt; }}

.signoff {{ margin-top: 28pt; padding-top: 18pt; border-top: 0.75pt solid {DNSF_BLACK}; }}
.signoff-text {{ font-family: 'Inter','Arial',sans-serif; font-style: italic; font-size: 12pt; color: {DNSF_BLACK}; }}
.signoff-name {{ font-family: 'Montserrat','Arial',sans-serif; font-size: 16pt; font-weight: 700; color: {DNSF_BLUE}; margin-top: 4pt; }}

.figure {{ margin: 22pt 0; padding: 14pt 0 0 0; text-align: center; page-break-inside: avoid; }}
.figure svg {{ width: 100%; max-width: 5.8in; height: auto; }}
.figure figcaption {{ font-family: 'Inter','Arial',sans-serif; font-style: italic; font-size: 9pt; color: {TEXT_MUTED}; margin-top: 10pt; letter-spacing: 0.2pt; }}
"""


def render_pdf(out_path):
    doc = (
        '<!DOCTYPE html><html><head><meta charset="utf-8">'
        f'<title>{COVER_FIELDS["title"]}</title></head><body>'
        f'{_cover_html_pdf()}{BODY}</body></html>'
    )
    HTML(string=doc).write_pdf(str(out_path), stylesheets=[CSS(string=PDF_CSS)])


# ---------------------------------------------------------------------------
# RENDERER: Markdown (GitHub Flavored)
# ---------------------------------------------------------------------------

_MD_CALLOUT_LABEL = {"info": "NOTE", "warn": "WARNING", "secure": "TIP"}


def _blockquote(text):
    return "\n".join(f"> {line}" if line else ">" for line in text.split("\n"))


def render_md(out_path, components, figures_dir):
    title = COVER_FIELDS["title"].replace("<br>", " ").replace("<br/>", " ")
    parts = [f"# {title}", ""]

    if COVER_FIELDS.get("eyebrow"):
        parts += [f"_{html_to_plain(COVER_FIELDS['eyebrow'])}_", ""]
    if COVER_FIELDS.get("subtitle"):
        parts += [html_to_plain(COVER_FIELDS["subtitle"]), ""]
    if COVER_FIELDS.get("doc_type"):
        parts += [f"**{html_to_plain(COVER_FIELDS['doc_type'])}**", ""]
    if COVER_FIELDS.get("byline"):
        parts.append("| " + " | ".join(label for label, _ in COVER_FIELDS["byline"]) + " |")
        parts.append("|" + "|".join(["---"] * len(COVER_FIELDS["byline"])) + "|")
        parts.append("| " + " | ".join(value for _, value in COVER_FIELDS["byline"]) + " |")
        parts.append("")
    parts += ["---", ""]

    for c in components:
        kind = c[0]
        if kind == "section":
            parts += [f"## {c[1]}", ""]
        elif kind == "paragraph":
            parts += [html_to_md_inline(c[1]), ""]
        elif kind == "callout":
            _, variant, ctitle, cbody = c
            label = _MD_CALLOUT_LABEL.get(variant, "NOTE")
            block = f"[!{label}]\n"
            if ctitle:
                block += f"**{ctitle}**\n\n"
            block += html_to_md_inline(cbody)
            parts += [_blockquote(block), ""]
        elif kind == "step":
            _, num, stitle, sbody = c
            heading = f"### {num}. {stitle}" if num else f"### {stitle}"
            parts += [heading, "", html_to_md_inline(sbody), ""]
        elif kind == "prompt":
            _, label, body = c
            parts.append("```text")
            if label:
                parts.append(f"# {label}")
            parts.append(html_to_plain(body))
            parts += ["```", ""]
        elif kind == "tools":
            for name, meta, desc in c[1]:
                line = f"- **{name}**"
                if meta:
                    line += f" ({meta})"
                if desc:
                    line += f". {html_to_md_inline(desc)}"
                parts.append(line)
            parts.append("")
        elif kind == "checklist":
            for item in c[1]:
                parts.append(f"- [ ] {html_to_md_inline(item)}")
            parts.append("")
        elif kind == "signoff":
            _, text, name = c
            parts += ["---", "", f"_{text}_", "", f"**{name}**", ""]
        elif kind == "figure":
            _, idx, svg_str, caption = c
            png_name = f"fig-{idx:02d}.png"
            (figures_dir / png_name).write_bytes(rasterize_svg(svg_str))
            rel = f"{figures_dir.name}/{png_name}"
            parts += [f"![{caption}]({rel})", ""]
            if caption:
                parts += [f"_{caption}_", ""]

    out_path.write_text("\n".join(parts).rstrip() + "\n")


# ---------------------------------------------------------------------------
# RENDERER: DOCX (python-docx with brand styles)
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _add_left_border(cell, color, size=24):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:color"), color.lstrip("#"))
    borders.append(left)
    tc_pr.append(borders)


def _add_paragraph_bottom_border(p, hex_color, size=8):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_runs_from_html(paragraph, html, base_color=None, base_font="Inter",
                        base_size=11, force_color=False):
    soup = BeautifulSoup(html or "", "html.parser")

    def style(run, bold, italic, link):
        run.bold = bold
        run.italic = italic
        run.font.name = base_font
        run.font.size = Pt(base_size)
        if link:
            run.font.color.rgb = hex_to_rgb(DNSF_BLUE)
            run.font.underline = True
        elif base_color is not None and (force_color or run.font.color.rgb is None):
            run.font.color.rgb = base_color

    def walk(node, bold=False, italic=False, link=None):
        for child in node.children:
            if isinstance(child, NavigableString):
                txt = str(child)
                if not txt:
                    continue
                style(paragraph.add_run(txt), bold, italic, link)
                continue
            tag = child.name
            if tag in ("strong", "b"):
                walk(child, bold=True, italic=italic, link=link)
            elif tag in ("em", "i"):
                walk(child, bold=bold, italic=True, link=link)
            elif tag == "a":
                walk(child, bold=bold, italic=italic, link=child.get("href"))
            elif tag == "br":
                paragraph.add_run().add_break()
            else:
                walk(child, bold=bold, italic=italic, link=link)

    walk(soup)


def _para(container, text, *, font="Inter", size=11, bold=False,
          italic=False, color=None, align=None, space_after=6):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def _para_html(container, html, *, font="Inter", size=11, color=None,
               align=None, space_after=6, force_color=False):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    _add_runs_from_html(p, html, base_color=color, base_font=font,
                        base_size=size, force_color=force_color)
    return p


def _cover_docx(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    if COVER == "feature":
        band = doc.add_table(rows=1, cols=2)
        band.autofit = False
        band.columns[0].width = Cm(11)
        band.columns[1].width = Cm(7)
        cell_l = band.rows[0].cells[0]
        cell_r = band.rows[0].cells[1]
        for c in (cell_l, cell_r):
            _set_cell_shading(c, DNSF_BLUE)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        png = rasterize_svg(LOGO_WHITE_SVG, width_px=600)
        p = cell_l.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(io.BytesIO(png), width=Cm(5.5))
        _para(cell_r, COVER_FIELDS["doc_type"], font="Inter", size=11,
              italic=True, color=RGBColor(255, 255, 255),
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
        doc.add_paragraph()
    else:
        top = doc.add_table(rows=1, cols=2)
        top.autofit = False
        top.columns[0].width = Cm(11)
        top.columns[1].width = Cm(7)
        cell_l = top.rows[0].cells[0]
        cell_r = top.rows[0].cells[1]
        for c in (cell_l, cell_r):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        png = rasterize_svg(LOGO_DARK_SVG, width_px=600)
        p = cell_l.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(io.BytesIO(png), width=Cm(5.5))
        _para(cell_r, COVER_FIELDS["doc_type"], font="Inter", size=11,
              italic=True, color=hex_to_rgb(DNSF_BLACK),
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
        rule = doc.add_paragraph()
        rule.paragraph_format.space_after = Pt(0)
        _add_paragraph_bottom_border(rule, DNSF_BLACK, size=6)

    doc.add_paragraph()
    if COVER_FIELDS.get("eyebrow"):
        _para(doc, html_to_plain(COVER_FIELDS["eyebrow"]),
              font="Inter", size=13, italic=True,
              color=hex_to_rgb(DNSF_BLUE), space_after=12)

    title_text = COVER_FIELDS["title"].replace("<br>", "\n").replace("<br/>", "\n")
    for line in title_text.split("\n"):
        _para(doc, line, font="Montserrat", size=44, bold=True,
              color=hex_to_rgb(DNSF_BLACK), space_after=2)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    _add_paragraph_bottom_border(rule, DNSF_BLUE, size=24)
    rule_run = rule.add_run(" " * 6)
    rule_run.font.size = Pt(2)

    if COVER_FIELDS.get("subtitle"):
        _para(doc, html_to_plain(COVER_FIELDS["subtitle"]),
              font="Inter", size=14, color=hex_to_rgb(DNSF_BLACK),
              space_after=24)

    if COVER_FIELDS.get("byline"):
        byline = doc.add_table(rows=2, cols=len(COVER_FIELDS["byline"]))
        byline.autofit = True
        for i, (label, value) in enumerate(COVER_FIELDS["byline"]):
            _para(byline.rows[0].cells[i], label.upper(),
                  font="Montserrat", size=8, bold=True,
                  color=hex_to_rgb(TEXT_MUTED), space_after=2)
            _para(byline.rows[1].cells[i], value,
                  font="Inter", size=11,
                  color=hex_to_rgb(DNSF_BLACK), space_after=0)

    doc.add_page_break()


def _docx_callout(doc, variant, title, body_html):
    color = {"info": DNSF_BLUE, "warn": DNSF_MAGENTA, "secure": DNSF_SKY}[variant]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, BG_LIGHT)
    _add_left_border(cell, color)
    cell.paragraphs[0].text = ""
    if title:
        _para(cell, title.upper(), font="Montserrat", size=8, bold=True,
              color=hex_to_rgb(color), space_after=4)
    _para_html(cell, body_html, color=hex_to_rgb(DNSF_BLACK),
               space_after=0, force_color=True)
    doc.add_paragraph()


def _docx_step(doc, num, title, body_html):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.5)
    table.columns[1].width = Cm(15.5)
    num_cell = table.rows[0].cells[0]
    body_cell = table.rows[0].cells[1]
    _para(num_cell, str(num), font="Montserrat", size=18, bold=True,
          color=hex_to_rgb(DNSF_BLUE), space_after=0)
    _para(body_cell, title, font="Montserrat", size=12, bold=True,
          color=hex_to_rgb(DNSF_BLACK), space_after=2)
    _para_html(body_cell, body_html, color=hex_to_rgb(DNSF_BLACK),
               space_after=0, force_color=True)


def _docx_prompt(doc, label, body_html):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, DNSF_BLACK)
    _add_left_border(cell, DNSF_BLUE)
    cell.paragraphs[0].text = ""
    if label:
        _para(cell, label.upper(), font="Montserrat", size=8, bold=True,
              color=hex_to_rgb(DNSF_BLUE), space_after=4)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _add_runs_from_html(p, body_html,
                        base_color=RGBColor(255, 255, 255), force_color=True)
    for r in p.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


def render_docx(out_path, components, figures_dir):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Inter"
    normal.font.size = Pt(11)

    _cover_docx(doc)

    for c in components:
        kind = c[0]
        if kind == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(c[1])
            run.font.name = "Montserrat"
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = hex_to_rgb(DNSF_BLACK)
            _add_paragraph_bottom_border(p, DNSF_BLACK, size=8)
        elif kind == "paragraph":
            _para_html(doc, c[1], color=hex_to_rgb(DNSF_BLACK),
                       space_after=8, force_color=True)
        elif kind == "callout":
            _docx_callout(doc, c[1], c[2], c[3])
        elif kind == "step":
            _docx_step(doc, c[1], c[2], c[3])
        elif kind == "prompt":
            _docx_prompt(doc, c[1], c[2])
        elif kind == "tools":
            for name, meta, desc in c[1]:
                _para(doc, name, font="Montserrat", size=13, bold=True,
                      color=hex_to_rgb(DNSF_BLACK), space_after=2)
                if meta:
                    _para(doc, meta, font="Inter", size=10, italic=True,
                          color=hex_to_rgb(DNSF_BLUE), space_after=2)
                if desc:
                    _para_html(doc, desc, color=hex_to_rgb(DNSF_BLACK),
                               space_after=8, force_color=True)
        elif kind == "checklist":
            for i, item in enumerate(c[1], start=1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run_num = p.add_run(f"{i:02d}.  ")
                run_num.font.name = "Montserrat"
                run_num.font.size = Pt(11)
                run_num.bold = True
                run_num.font.color.rgb = hex_to_rgb(DNSF_BLUE)
                _add_runs_from_html(p, item,
                                    base_color=hex_to_rgb(DNSF_BLACK),
                                    force_color=True)
        elif kind == "signoff":
            _, stext, sname = c
            doc.add_paragraph()
            _para(doc, stext, font="Inter", size=12, italic=True,
                  color=hex_to_rgb(DNSF_BLACK), space_after=2)
            _para(doc, sname, font="Montserrat", size=16, bold=True,
                  color=hex_to_rgb(DNSF_BLUE), space_after=8)
        elif kind == "figure":
            _, idx, svg_str, caption = c
            png_bytes = rasterize_svg(svg_str, width_px=1200)
            (figures_dir / f"fig-{idx:02d}.png").write_bytes(png_bytes)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(png_bytes), width=Cm(14))
            if caption:
                _para(doc, caption, font="Inter", size=9, italic=True,
                      color=hex_to_rgb(TEXT_MUTED),
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# RENDERER: Confluence storage format (XHTML)
# ---------------------------------------------------------------------------

_CONFLUENCE_MACRO = {"info": "info", "warn": "warning", "secure": "tip"}


def _conf_inline(html):
    return (html or "").strip()


def render_confluence(out_path, components, figures_dir):
    parts = []
    title = COVER_FIELDS["title"].replace("<br>", " ").replace("<br/>", " ")
    parts.append(f"<h1>{html_escape(title)}</h1>")
    if COVER_FIELDS.get("eyebrow"):
        parts.append(f'<p><em>{html_escape(html_to_plain(COVER_FIELDS["eyebrow"]))}</em></p>')
    if COVER_FIELDS.get("subtitle"):
        parts.append(f"<p>{html_escape(html_to_plain(COVER_FIELDS['subtitle']))}</p>")
    if COVER_FIELDS.get("byline"):
        rows = "".join(
            f"<tr><th>{html_escape(label)}</th><td>{html_escape(value)}</td></tr>"
            for label, value in COVER_FIELDS["byline"]
        )
        parts.append(f"<table><tbody>{rows}</tbody></table>")
    parts.append("<hr/>")

    for c in components:
        kind = c[0]
        if kind == "section":
            parts.append(f"<h2>{html_escape(c[1])}</h2>")
        elif kind == "paragraph":
            parts.append(f"<p>{_conf_inline(c[1])}</p>")
        elif kind == "callout":
            _, variant, ctitle, cbody = c
            macro = _CONFLUENCE_MACRO[variant]
            inner = ""
            if ctitle:
                inner += f"<p><strong>{html_escape(ctitle)}</strong></p>"
            inner += f"<p>{_conf_inline(cbody)}</p>"
            parts.append(
                f'<ac:structured-macro ac:name="{macro}" ac:schema-version="1">'
                f"<ac:rich-text-body>{inner}</ac:rich-text-body>"
                f"</ac:structured-macro>"
            )
        elif kind == "step":
            _, num, stitle, sbody = c
            heading = f"{num}. {stitle}" if num else stitle
            parts.append(f"<h3>{html_escape(heading)}</h3>")
            parts.append(f"<p>{_conf_inline(sbody)}</p>")
        elif kind == "prompt":
            _, label, body = c
            cdata = html_to_plain(body)
            macro = '<ac:structured-macro ac:name="code" ac:schema-version="1">'
            if label:
                macro += f'<ac:parameter ac:name="title">{html_escape(label)}</ac:parameter>'
            macro += (
                f"<ac:plain-text-body><![CDATA[{cdata}]]></ac:plain-text-body>"
                "</ac:structured-macro>"
            )
            parts.append(macro)
        elif kind == "tools":
            rows = "".join(
                f"<tr><td><strong>{html_escape(name)}</strong></td>"
                f"<td><em>{html_escape(meta)}</em></td>"
                f"<td>{_conf_inline(desc)}</td></tr>"
                for name, meta, desc in c[1]
            )
            parts.append(
                "<table><thead><tr><th>Tool</th><th>Type</th><th>Why</th></tr>"
                f"</thead><tbody>{rows}</tbody></table>"
            )
        elif kind == "checklist":
            items = "".join(
                f"<ac:task><ac:task-status>incomplete</ac:task-status>"
                f"<ac:task-body>{_conf_inline(item)}</ac:task-body></ac:task>"
                for item in c[1]
            )
            parts.append(f"<ac:task-list>{items}</ac:task-list>")
        elif kind == "signoff":
            _, stext, sname = c
            parts.append(
                f"<hr/><p><em>{html_escape(stext)}</em></p>"
                f"<p><strong>{html_escape(sname)}</strong></p>"
            )
        elif kind == "figure":
            _, idx, svg_str, caption = c
            png_name = f"fig-{idx:02d}.png"
            (figures_dir / png_name).write_bytes(rasterize_svg(svg_str, width_px=1200))
            parts.append(
                f'<ac:image ac:align="center" ac:width="600">'
                f'<ri:attachment ri:filename="{png_name}" />'
                f"</ac:image>"
            )
            if caption:
                parts.append(
                    f'<p style="text-align: center;"><em>{html_escape(caption)}</em></p>'
                )

    header = (
        "<!-- DNSFilter guide, Confluence storage format. "
        "Push via REST API (POST /wiki/rest/api/content) or paste into a "
        "Confluence page in source view. Upload the PNGs in the figures "
        "folder as page attachments so the figure references resolve. -->\n"
    )
    out_path.write_text(header + "\n".join(parts) + "\n")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    components = parse_body(BODY)
    figures_dir = OUT_DIR / f"{OUT_BASENAME}-figures"
    if any(c[0] == "figure" for c in components) and (
        "md" in FORMATS or "docx" in FORMATS or "confluence" in FORMATS
    ):
        figures_dir.mkdir(parents=True, exist_ok=True)

    written = []
    if "pdf" in FORMATS:
        path = OUT_DIR / f"{OUT_BASENAME}.pdf"
        render_pdf(path)
        written.append(path)
    if "md" in FORMATS:
        path = OUT_DIR / f"{OUT_BASENAME}.md"
        render_md(path, components, figures_dir)
        written.append(path)
    if "docx" in FORMATS:
        path = OUT_DIR / f"{OUT_BASENAME}.docx"
        render_docx(path, components, figures_dir)
        written.append(path)
    if "confluence" in FORMATS:
        path = OUT_DIR / f"{OUT_BASENAME}.confluence.xhtml"
        render_confluence(path, components, figures_dir)
        written.append(path)

    for p in written:
        print(f"wrote {p}")


if __name__ == "__main__":
    main()
